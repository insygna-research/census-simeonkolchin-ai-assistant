"""
Outline Tool - Integration with Outline Wiki API
"""

import httpx
import io
import re
import logging
from typing import Dict, List, Any, Optional

logger = logging.getLogger(__name__)


class OutlineTool:
    """Tool for interacting with Outline Wiki"""
    
    def __init__(self, url: str, api_token: str):
        """
        Initialize Outline tool
        
        Args:
            url: Outline Wiki URL (e.g., https://outline.example.com)
            api_token: Outline API token
        """
        self.base_url = url.rstrip('/')
        self.api_token = api_token
        self.headers = {
            'Authorization': f'Bearer {self.api_token}',
            'Content-Type': 'application/json'
        }
    
    def _post(self, endpoint: str, data: dict = None) -> dict:
        """Make a POST request to Outline API"""
        url = f"{self.base_url}/api/{endpoint}"
        try:
            response = httpx.post(
                url,
                headers=self.headers,
                json=data or {},
                timeout=30.0
            )
            response.raise_for_status()
            return response.json()
        except httpx.HTTPStatusError as e:
            # Include response body for better debugging
            error_detail = str(e)
            try:
                error_body = e.response.json()
                error_detail = f"{e}: {error_body}"
            except:
                error_detail = f"{e}: {e.response.text[:500]}"
            logger.error(f"Outline API error: {error_detail}")
            return {"ok": False, "error": error_detail}
        except httpx.HTTPError as e:
            return {"ok": False, "error": str(e)}
    
    async def _apost(self, endpoint: str, data: dict = None) -> dict:
        """Make an async POST request to Outline API"""
        url = f"{self.base_url}/api/{endpoint}"
        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    url,
                    headers=self.headers,
                    json=data or {},
                    timeout=30.0
                )
                response.raise_for_status()
                return response.json()
        except httpx.HTTPStatusError as e:
            # Include response body for better debugging
            error_detail = str(e)
            try:
                error_body = e.response.json()
                error_detail = f"{e}: {error_body}"
            except:
                error_detail = f"{e}: {e.response.text[:500]}"
            logger.error(f"Outline API error: {error_detail}")
            return {"ok": False, "error": error_detail}
        except httpx.HTTPError as e:
            return {"ok": False, "error": str(e)}
    
    # ============ Search & Retrieval ============
    
    def search_documents(self, query: str, limit: int = 10) -> Dict[str, Any]:
        """
        Search documents in Outline
        
        Args:
            query: Search query
            limit: Maximum number of results
        
        Returns:
            Dictionary with search results
        """
        result = self._post('documents.search', {'query': query, 'limit': limit})
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Search failed')}
        
        # Format results for better readability
        documents = []
        for item in result.get('data', []):
            doc = item.get('document', {})
            documents.append({
                'id': doc.get('id'),
                'title': doc.get('title'),
                'url': doc.get('url'),
                'context': item.get('context', ''),
                'collection_id': doc.get('collectionId'),
                'updated_at': doc.get('updatedAt')
            })
        
        return {
            "success": True,
            "query": query,
            "count": len(documents),
            "documents": documents
        }
    
    def get_document(self, document_id: str) -> Dict[str, Any]:
        """
        Get a document by ID
        
        Args:
            document_id: Document ID
        
        Returns:
            Dictionary with document data
        """
        result = self._post('documents.info', {'id': document_id})
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Document not found')}
        
        doc = result.get('data', {})
        return {
            "success": True,
            "id": doc.get('id'),
            "title": doc.get('title'),
            "text": doc.get('text'),
            "url": doc.get('url'),
            "collection_id": doc.get('collectionId'),
            "parent_document_id": doc.get('parentDocumentId'),
            "created_at": doc.get('createdAt'),
            "updated_at": doc.get('updatedAt'),
            "created_by": doc.get('createdBy', {}).get('name')
        }
    
    def list_collections(self) -> Dict[str, Any]:
        """
        List all collections
        
        Returns:
            Dictionary with collections list
        """
        result = self._post('collections.list', {'limit': 100})
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Failed to list collections')}
        
        collections = []
        for coll in result.get('data', []):
            collections.append({
                'id': coll.get('id'),
                'name': coll.get('name'),
                'description': coll.get('description'),
                'icon': coll.get('icon'),
                'color': coll.get('color'),
                'url': coll.get('url')
            })
        
        return {
            "success": True,
            "count": len(collections),
            "collections": collections
        }
    
    def list_documents(self, collection_id: Optional[str] = None, limit: int = 50) -> Dict[str, Any]:
        """
        List documents, optionally filtered by collection
        
        Args:
            collection_id: Optional collection ID to filter by
            limit: Maximum number of documents
        
        Returns:
            Dictionary with documents list
        """
        data = {'limit': limit}
        if collection_id:
            data['collectionId'] = collection_id
        
        result = self._post('documents.list', data)
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Failed to list documents')}
        
        documents = []
        for doc in result.get('data', []):
            documents.append({
                'id': doc.get('id'),
                'title': doc.get('title'),
                'url': doc.get('url'),
                'collection_id': doc.get('collectionId'),
                'updated_at': doc.get('updatedAt')
            })
        
        return {
            "success": True,
            "count": len(documents),
            "documents": documents
        }
    
    # ============ Create & Update ============
    
    def create_document(
        self,
        title: str,
        text: str,
        collection_id: str,
        parent_document_id: Optional[str] = None,
        publish: bool = True
    ) -> Dict[str, Any]:
        """
        Create a new document in Outline
        
        Args:
            title: Document title
            text: Document content in Markdown
            collection_id: ID of the collection to create document in
            parent_document_id: Optional parent document ID for nested documents
            publish: Whether to publish the document immediately
        
        Returns:
            Dictionary with created document data
        """
        # Validate inputs
        if not title or not title.strip():
            return {"error": "Title is required and cannot be empty"}
        if not collection_id or not collection_id.strip():
            return {"error": "collection_id is required and cannot be empty"}
        if text is None:
            text = ""
        
        data = {
            'title': title.strip(),
            'text': text,
            'collectionId': collection_id.strip(),
            'publish': publish
        }
        
        if parent_document_id and parent_document_id.strip():
            data['parentDocumentId'] = parent_document_id.strip()
        
        logger.info(f"Creating document: title='{title}', collection_id='{collection_id}', text_length={len(text)}")
        
        result = self._post('documents.create', data)
        
        if not result.get('ok'):
            error_msg = result.get('error', 'Failed to create document')
            logger.error(f"Failed to create document: {error_msg}")
            return {"error": error_msg}
        
        doc = result.get('data', {})
        return {
            "success": True,
            "id": doc.get('id'),
            "title": doc.get('title'),
            "url": doc.get('url'),
            "collection_id": doc.get('collectionId'),
            "message": f"Документ '{title}' успешно создан"
        }
    
    def update_document(
        self,
        document_id: str,
        title: Optional[str] = None,
        text: Optional[str] = None,
        append: bool = False,
        done: bool = False
    ) -> Dict[str, Any]:
        """
        Update an existing document
        
        Args:
            document_id: ID of document to update
            title: New title (optional)
            text: New text content (optional)
            append: If True, append text instead of replacing
            done: Mark editing as done
        
        Returns:
            Dictionary with update result
        """
        data = {
            'id': document_id,
            'done': done
        }
        
        if title:
            data['title'] = title
        
        if text:
            data['text'] = text
            if append:
                data['append'] = True
        
        result = self._post('documents.update', data)
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Failed to update document')}
        
        doc = result.get('data', {})
        return {
            "success": True,
            "id": doc.get('id'),
            "title": doc.get('title'),
            "url": doc.get('url'),
            "message": f"Документ успешно обновлён"
        }
    
    # ============ Attachments & Document Reading ============

    def _extract_attachment_urls(self, text: str) -> List[Dict[str, str]]:
        """Extract attachment URLs from Outline document markdown text"""
        attachments = []
        # Pattern 1: markdown links like [filename.docx](/api/attachments.redirect?id=UUID)
        pattern1 = r'\[([^\]]+)\]\((/api/attachments\.redirect\?id=[a-f0-9-]+)\)'
        for match in re.finditer(pattern1, text):
            attachments.append({
                'name': match.group(1),
                'url': match.group(2),
            })
        # Pattern 2: bare attachment URLs /api/attachments/UUID/filename
        pattern2 = r'(/api/attachments/[a-f0-9-]+/[^\s\)]+)'
        for match in re.finditer(pattern2, text):
            url = match.group(1)
            name = url.split('/')[-1]
            # Avoid duplicates
            if not any(a['url'] == url for a in attachments):
                attachments.append({
                    'name': name,
                    'url': url,
                })
        return attachments

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file bytes"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            return '\n\n'.join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return f"[Ошибка чтения DOCX: {str(e)}]"

    def _extract_text_from_doc(self, content: bytes) -> str:
        """
        Extract text from legacy DOC file bytes.
        Uses a simple binary text extraction since python-docx only supports DOCX.
        """
        try:
            # Try as DOCX first (some .doc files are actually DOCX)
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                return '\n\n'.join(paragraphs)
        except Exception:
            pass

        # Fallback: extract readable text from binary DOC
        try:
            text_parts = []
            # DOC files contain text between certain markers
            # Simple approach: find UTF-8/Latin-1 text sequences
            decoded = content.decode('utf-8', errors='ignore')
            # Filter out binary noise - keep lines with mostly printable chars
            for line in decoded.split('\n'):
                clean = line.strip()
                if len(clean) < 3:
                    continue
                # Count printable vs non-printable
                printable = sum(1 for c in clean if c.isprintable() or c in '\t\r\n')
                if len(clean) > 0 and printable / len(clean) > 0.8:
                    text_parts.append(clean)

            if text_parts:
                return '\n'.join(text_parts)
            
            # Try cp1251 (common for Russian docs)
            decoded = content.decode('cp1251', errors='ignore')
            for line in decoded.split('\n'):
                clean = line.strip()
                if len(clean) < 3:
                    continue
                printable = sum(1 for c in clean if c.isprintable() or c in '\t\r\n')
                if len(clean) > 0 and printable / len(clean) > 0.8:
                    text_parts.append(clean)
            
            if text_parts:
                return '\n'.join(text_parts)
            
            return "[Не удалось извлечь текст из DOC файла. Попробуйте конвертировать в DOCX.]"
        except Exception as e:
            return f"[Ошибка чтения DOC: {str(e)}]"

    def get_document_attachments(self, document_id: str) -> Dict[str, Any]:
        """
        Get list of attachments from an Outline document.
        Parses attachment links from document markdown.
        
        Args:
            document_id: Document ID
        
        Returns:
            Dictionary with attachment info
        """
        # First get the document
        doc_result = self.get_document(document_id)
        if 'error' in doc_result:
            return doc_result
        
        text = doc_result.get('text', '')
        attachments = self._extract_attachment_urls(text)
        
        return {
            "success": True,
            "document_id": document_id,
            "document_title": doc_result.get('title', ''),
            "count": len(attachments),
            "attachments": attachments
        }

    def read_document_attachment(self, document_id: str, attachment_name: Optional[str] = None) -> Dict[str, Any]:
        """
        Download and read text from an attachment in an Outline document.
        Supports DOCX, DOC, and TXT files.
        
        Args:
            document_id: Document ID containing the attachment
            attachment_name: Optional filename to filter (if multiple attachments)
        
        Returns:
            Dictionary with extracted text
        """
        # Get attachments list
        att_result = self.get_document_attachments(document_id)
        if 'error' in att_result:
            return att_result
        
        attachments = att_result.get('attachments', [])
        if not attachments:
            return {"error": "В документе нет вложений"}
        
        # Filter by name if specified
        target = None
        if attachment_name:
            for att in attachments:
                if attachment_name.lower() in att['name'].lower():
                    target = att
                    break
            if not target:
                return {
                    "error": f"Вложение '{attachment_name}' не найдено",
                    "available": [a['name'] for a in attachments]
                }
        else:
            # If multiple, read the first one; if only one, read it
            target = attachments[0]
        
        # Download the attachment
        try:
            url = f"{self.base_url}{target['url']}"
            response = httpx.get(
                url,
                headers={'Authorization': f'Bearer {self.api_token}'},
                follow_redirects=True,
                timeout=60.0
            )
            response.raise_for_status()
            content = response.content
            filename = target['name'].lower()
            
            logger.info(f"Downloaded attachment: {target['name']} ({len(content)} bytes)")
            
            # Extract text based on file type
            if filename.endswith('.docx'):
                extracted_text = self._extract_text_from_docx(content)
            elif filename.endswith('.doc'):
                extracted_text = self._extract_text_from_doc(content)
            elif filename.endswith('.txt') or filename.endswith('.md'):
                # Try UTF-8, then cp1251
                try:
                    extracted_text = content.decode('utf-8')
                except UnicodeDecodeError:
                    extracted_text = content.decode('cp1251', errors='replace')
            else:
                # Try to read as text anyway
                try:
                    extracted_text = content.decode('utf-8')
                except UnicodeDecodeError:
                    extracted_text = f"[Неподдерживаемый формат файла: {target['name']}]"
            
            # No truncation - return full text (context management is handled by agent)
            truncated = False
            
            return {
                "success": True,
                "filename": target['name'],
                "document_id": document_id,
                "document_title": att_result.get('document_title', ''),
                "size_bytes": len(content),
                "text_length": len(extracted_text),
                "truncated": truncated,
                "text": extracted_text
            }
            
        except httpx.HTTPError as e:
            return {"error": f"Ошибка скачивания вложения: {str(e)}"}
        except Exception as e:
            return {"error": f"Ошибка обработки вложения: {str(e)}"}

    async def aget_document_attachments(self, document_id: str) -> Dict[str, Any]:
        """Async version of get_document_attachments"""
        doc_result = await self.aget_document(document_id)
        if 'error' in doc_result:
            return doc_result
        
        text = doc_result.get('text', '')
        attachments = self._extract_attachment_urls(text)
        
        return {
            "success": True,
            "document_id": document_id,
            "document_title": doc_result.get('title', ''),
            "count": len(attachments),
            "attachments": attachments
        }

    async def aread_document_attachment(self, document_id: str, attachment_name: Optional[str] = None) -> Dict[str, Any]:
        """Async version of read_document_attachment"""
        # Get attachments list
        att_result = await self.aget_document_attachments(document_id)
        if 'error' in att_result:
            return att_result
        
        attachments = att_result.get('attachments', [])
        if not attachments:
            return {"error": "В документе нет вложений"}
        
        # Filter by name if specified
        target = None
        if attachment_name:
            for att in attachments:
                if attachment_name.lower() in att['name'].lower():
                    target = att
                    break
            if not target:
                return {
                    "error": f"Вложение '{attachment_name}' не найдено",
                    "available": [a['name'] for a in attachments]
                }
        else:
            target = attachments[0]
        
        # Download the attachment
        try:
            url = f"{self.base_url}{target['url']}"
            async with httpx.AsyncClient() as client:
                response = await client.get(
                    url,
                    headers={'Authorization': f'Bearer {self.api_token}'},
                    follow_redirects=True,
                    timeout=60.0
                )
                response.raise_for_status()
                content = response.content
            
            filename = target['name'].lower()
            logger.info(f"Downloaded attachment: {target['name']} ({len(content)} bytes)")
            
            # Extract text based on file type
            if filename.endswith('.docx'):
                extracted_text = self._extract_text_from_docx(content)
            elif filename.endswith('.doc'):
                extracted_text = self._extract_text_from_doc(content)
            elif filename.endswith('.txt') or filename.endswith('.md'):
                try:
                    extracted_text = content.decode('utf-8')
                except UnicodeDecodeError:
                    extracted_text = content.decode('cp1251', errors='replace')
            else:
                try:
                    extracted_text = content.decode('utf-8')
                except UnicodeDecodeError:
                    extracted_text = f"[Неподдерживаемый формат файла: {target['name']}]"
            
            # No truncation - return full text (context management is handled by agent)
            truncated = False
            
            return {
                "success": True,
                "filename": target['name'],
                "document_id": document_id,
                "document_title": att_result.get('document_title', ''),
                "size_bytes": len(content),
                "text_length": len(extracted_text),
                "truncated": truncated,
                "text": extracted_text
            }
            
        except httpx.HTTPError as e:
            return {"error": f"Ошибка скачивания вложения: {str(e)}"}
        except Exception as e:
            return {"error": f"Ошибка обработки вложения: {str(e)}"}

    # ============ Async versions ============
    
    async def asearch_documents(self, query: str, limit: int = 10) -> Dict[str, Any]:
        """Async version of search_documents"""
        result = await self._apost('documents.search', {'query': query, 'limit': limit})
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Search failed')}
        
        documents = []
        for item in result.get('data', []):
            doc = item.get('document', {})
            documents.append({
                'id': doc.get('id'),
                'title': doc.get('title'),
                'url': doc.get('url'),
                'context': item.get('context', ''),
                'collection_id': doc.get('collectionId'),
                'updated_at': doc.get('updatedAt')
            })
        
        return {
            "success": True,
            "query": query,
            "count": len(documents),
            "documents": documents
        }
    
    async def aget_document(self, document_id: str) -> Dict[str, Any]:
        """Async version of get_document"""
        result = await self._apost('documents.info', {'id': document_id})
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Document not found')}
        
        doc = result.get('data', {})
        return {
            "success": True,
            "id": doc.get('id'),
            "title": doc.get('title'),
            "text": doc.get('text'),
            "url": doc.get('url'),
            "collection_id": doc.get('collectionId'),
            "parent_document_id": doc.get('parentDocumentId'),
            "created_at": doc.get('createdAt'),
            "updated_at": doc.get('updatedAt'),
            "created_by": doc.get('createdBy', {}).get('name')
        }
    
    async def alist_collections(self) -> Dict[str, Any]:
        """Async version of list_collections"""
        result = await self._apost('collections.list', {'limit': 100})
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Failed to list collections')}
        
        collections = []
        for coll in result.get('data', []):
            collections.append({
                'id': coll.get('id'),
                'name': coll.get('name'),
                'description': coll.get('description'),
                'icon': coll.get('icon'),
                'color': coll.get('color'),
                'url': coll.get('url')
            })
        
        return {
            "success": True,
            "count": len(collections),
            "collections": collections
        }
    
    async def acreate_document(
        self,
        title: str,
        text: str,
        collection_id: str,
        parent_document_id: Optional[str] = None,
        publish: bool = True
    ) -> Dict[str, Any]:
        """Async version of create_document"""
        # Validate inputs
        if not title or not title.strip():
            return {"error": "Title is required and cannot be empty"}
        if not collection_id or not collection_id.strip():
            return {"error": "collection_id is required and cannot be empty"}
        if text is None:
            text = ""
        
        data = {
            'title': title.strip(),
            'text': text,
            'collectionId': collection_id.strip(),
            'publish': publish
        }
        
        if parent_document_id and parent_document_id.strip():
            data['parentDocumentId'] = parent_document_id.strip()
        
        logger.info(f"Creating document: title='{title}', collection_id='{collection_id}', text_length={len(text)}")
        
        result = await self._apost('documents.create', data)
        
        if not result.get('ok'):
            error_msg = result.get('error', 'Failed to create document')
            logger.error(f"Failed to create document: {error_msg}")
            return {"error": error_msg}
        
        doc = result.get('data', {})
        return {
            "success": True,
            "id": doc.get('id'),
            "title": doc.get('title'),
            "url": doc.get('url'),
            "collection_id": doc.get('collectionId'),
            "message": f"Документ '{title}' успешно создан"
        }
    
    async def aupdate_document(
        self,
        document_id: str,
        title: Optional[str] = None,
        text: Optional[str] = None,
        append: bool = False,
        done: bool = False
    ) -> Dict[str, Any]:
        """Async version of update_document"""
        data = {
            'id': document_id,
            'done': done
        }
        
        if title:
            data['title'] = title
        
        if text:
            data['text'] = text
            if append:
                data['append'] = True
        
        result = await self._apost('documents.update', data)
        
        if not result.get('ok'):
            return {"error": result.get('error', 'Failed to update document')}
        
        doc = result.get('data', {})
        return {
            "success": True,
            "id": doc.get('id'),
            "title": doc.get('title'),
            "url": doc.get('url'),
            "message": f"Документ успешно обновлён"
        }


# Tool definitions for LLM
OUTLINE_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "outline_search",
            "description": "Поиск информации в базе знаний Outline Wiki. Используй для поиска документов по ключевым словам.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Поисковый запрос"
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Максимальное количество результатов (по умолчанию 10)",
                        "default": 10
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "outline_get_document",
            "description": "Получить полный текст документа из Outline по его ID. Используй после поиска для получения полного содержимого.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "string",
                        "description": "ID документа в Outline"
                    }
                },
                "required": ["document_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "outline_list_collections",
            "description": "Получить список всех коллекций (категорий) в Outline. Используй для навигации по структуре базы знаний.",
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "outline_create_document",
            "description": "Создать новый документ в Outline. Используй для сохранения аналитических отчётов, выводов, рекомендаций.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "Заголовок документа"
                    },
                    "text": {
                        "type": "string",
                        "description": "Содержимое документа в формате Markdown"
                    },
                    "collection_id": {
                        "type": "string",
                        "description": "ID коллекции, в которую добавить документ"
                    },
                    "parent_document_id": {
                        "type": "string",
                        "description": "Опционально: ID родительского документа для создания вложенности"
                    }
                },
                "required": ["title", "text", "collection_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "outline_update_document",
            "description": "Обновить существующий документ в Outline. Используй для добавления новой информации или исправления.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "string",
                        "description": "ID документа для обновления"
                    },
                    "title": {
                        "type": "string",
                        "description": "Новый заголовок (опционально)"
                    },
                    "text": {
                        "type": "string",
                        "description": "Новое содержимое или дополнительный текст"
                    },
                    "append": {
                        "type": "boolean",
                        "description": "Если true, добавить текст в конец вместо замены",
                        "default": False
                    }
                },
                "required": ["document_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "outline_get_attachments",
            "description": "Получить список вложений (файлов) прикреплённых к документу Outline. Используй чтобы узнать какие файлы (DOC, DOCX, PDF и др.) прикреплены к странице.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "string",
                        "description": "ID документа в Outline"
                    }
                },
                "required": ["document_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "outline_read_attachment",
            "description": "Скачать и прочитать текст из вложенного файла в документе Outline. Поддерживает DOC, DOCX, TXT файлы. Используй для чтения транскрипций и других текстовых документов прикреплённых к страницам Outline.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "string",
                        "description": "ID документа в Outline, содержащего вложение"
                    },
                    "attachment_name": {
                        "type": "string",
                        "description": "Имя файла вложения (опционально, если не указано - читает первое вложение)"
                    }
                },
                "required": ["document_id"]
            }
        }
    }
]


class MultiOutlineTool:
    """Wraps multiple OutlineTool instances for multi-server support."""

    def __init__(self, servers: List[Dict[str, str]]):
        """
        Args:
            servers: list of {name, url, api_token, category}
        """
        self._tools: List[Dict[str, Any]] = []
        for srv in servers:
            tool = OutlineTool(url=srv["url"], api_token=srv["api_token"])
            self._tools.append({
                "name": srv["name"],
                "category": srv.get("category", "internal"),
                "url": srv["url"],
                "tool": tool,
            })
        logger.info(f"MultiOutlineTool initialized with {len(self._tools)} server(s): "
                     + ", ".join(f'{t["name"]} ({t["category"]})' for t in self._tools))

    def _label(self, t: dict) -> str:
        return f'[{t["name"]} ({t["category"]})]'

    def _dedup(self, items: list, key: str = "id") -> list:
        seen = set()
        result = []
        for item in items:
            k = item.get(key)
            if k and k not in seen:
                seen.add(k)
                result.append(item)
        return result

    # ── Sync (aggregating) ──

    def search_documents(self, query: str, limit: int = 10) -> Dict[str, Any]:
        all_docs = []
        for t in self._tools:
            result = t["tool"].search_documents(query, limit=limit)
            if "documents" in result:
                for doc in result["documents"]:
                    doc["_source"] = self._label(t)
                all_docs.extend(result["documents"])
        return {"documents": self._dedup(all_docs)[:limit]}

    def list_collections(self) -> Dict[str, Any]:
        all_cols = []
        for t in self._tools:
            result = t["tool"].list_collections()
            if "collections" in result:
                for col in result["collections"]:
                    col["_source"] = self._label(t)
                all_cols.extend(result["collections"])
        return {"collections": all_cols}

    # ── Sync (try each) ──

    def get_document(self, document_id: str) -> Dict[str, Any]:
        for t in self._tools:
            result = t["tool"].get_document(document_id)
            if "error" not in result or "ok" in result and result["ok"] is not False:
                if isinstance(result, dict):
                    result["_source"] = self._label(t)
                return result
        return {"error": f"Document {document_id} not found in any Outline server"}

    def create_document(self, **kwargs) -> Dict[str, Any]:
        for t in self._tools:
            result = t["tool"].create_document(**kwargs)
            if "error" not in result:
                result["_source"] = self._label(t)
                return result
        return {"error": "Failed to create document on any Outline server"}

    def update_document(self, **kwargs) -> Dict[str, Any]:
        for t in self._tools:
            result = t["tool"].update_document(**kwargs)
            if "error" not in result:
                result["_source"] = self._label(t)
                return result
        return {"error": "Failed to update document on any Outline server"}

    def get_document_attachments(self, document_id: str) -> Dict[str, Any]:
        for t in self._tools:
            result = t["tool"].get_document_attachments(document_id)
            if "error" not in result:
                return result
        return {"error": f"Attachments for {document_id} not found"}

    def read_document_attachment(self, document_id: str, attachment_name: Optional[str] = None) -> Dict[str, Any]:
        for t in self._tools:
            result = t["tool"].read_document_attachment(document_id, attachment_name)
            if "error" not in result:
                return result
        return {"error": f"Attachment not found in any server"}

    # ── Async (aggregating) ──

    async def asearch_documents(self, query: str, limit: int = 10) -> Dict[str, Any]:
        all_docs = []
        for t in self._tools:
            result = await t["tool"].asearch_documents(query, limit=limit)
            if "documents" in result:
                for doc in result["documents"]:
                    doc["_source"] = self._label(t)
                all_docs.extend(result["documents"])
        return {"documents": self._dedup(all_docs)[:limit]}

    async def alist_collections(self) -> Dict[str, Any]:
        all_cols = []
        for t in self._tools:
            result = await t["tool"].alist_collections()
            if "collections" in result:
                for col in result["collections"]:
                    col["_source"] = self._label(t)
                all_cols.extend(result["collections"])
        return {"collections": all_cols}

    # ── Async (try each) ──

    async def aget_document(self, document_id: str) -> Dict[str, Any]:
        for t in self._tools:
            result = await t["tool"].aget_document(document_id)
            if "error" not in result:
                result["_source"] = self._label(t)
                return result
        return {"error": f"Document {document_id} not found in any Outline server"}

    async def acreate_document(self, **kwargs) -> Dict[str, Any]:
        for t in self._tools:
            result = await t["tool"].acreate_document(**kwargs)
            if "error" not in result:
                result["_source"] = self._label(t)
                return result
        return {"error": "Failed to create document on any Outline server"}

    async def aupdate_document(self, **kwargs) -> Dict[str, Any]:
        for t in self._tools:
            result = await t["tool"].aupdate_document(**kwargs)
            if "error" not in result:
                result["_source"] = self._label(t)
                return result
        return {"error": "Failed to update document on any Outline server"}

    async def aget_document_attachments(self, document_id: str) -> Dict[str, Any]:
        for t in self._tools:
            result = await t["tool"].aget_document_attachments(document_id)
            if "error" not in result:
                return result
        return {"error": f"Attachments for {document_id} not found"}

    async def aread_document_attachment(self, document_id: str, attachment_name: Optional[str] = None) -> Dict[str, Any]:
        for t in self._tools:
            result = await t["tool"].aread_document_attachment(document_id, attachment_name)
            if "error" not in result:
                return result
        return {"error": f"Attachment not found in any server"}
